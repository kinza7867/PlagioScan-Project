from flask import Flask, render_template, request, jsonify, session
from checker import check_plagiarism
import hashlib
import json
import os
from datetime import datetime
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
import io
import traceback

app = Flask(__name__)
app.secret_key = 'your-secret-key-here-change-in-production'
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max
app.config['UPLOAD_EXTENSIONS'] = ['txt', 'pdf', 'docx', 'png', 'jpg', 'jpeg']
app.config['PERMANENT_SESSION_LIFETIME'] = 86400  # 24 hours

# Create upload folder if not exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'png', 'jpg', 'jpeg'}

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(filepath):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(filepath, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num} ---\n{page_text}"
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""

def extract_text_from_docx(filepath):
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        return text
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_image(filepath):
    """Extract text from image using OCR"""
    text = ""
    try:
        image = Image.open(filepath)
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        # Apply OCR with language support
        text = pytesseract.image_to_string(image, lang='eng')
        return text.strip()
    except Exception as e:
        print(f"OCR extraction error: {e}")
        return ""

def extract_text_from_file(filepath, filename):
    """Extract text from uploaded files based on extension"""
    ext = filename.rsplit('.', 1)[1].lower()
    text = ""
    
    try:
        if ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        
        elif ext == 'pdf':
            text = extract_text_from_pdf(filepath)
        
        elif ext == 'docx':
            text = extract_text_from_docx(filepath)
        
        elif ext in ['png', 'jpg', 'jpeg']:
            text = extract_text_from_image(filepath)
        
        # Clean up text
        if text:
            # Remove extra whitespace
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
        
        return text
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        traceback.print_exc()
        return ""

@app.route('/')
def index():
    """Home page"""
    return render_template('index.html')

@app.route('/features')
def features():
    """Features page"""
    return render_template('index.html')  # Single page app

@app.route('/pricing')
def pricing():
    """Pricing page"""
    return render_template('index.html')

@app.route('/api-docs')
def api_docs():
    """API documentation page"""
    return render_template('index.html')

@app.route('/about')
def about():
    """About page"""
    return render_template('index.html')

@app.route('/support')
def support():
    """Support page"""
    return render_template('index.html')

@app.route('/dashboard')
def dashboard():
    """Dashboard page"""
    return render_template('index.html')

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file uploads for text extraction"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'error': f'File type not supported. Allowed: {", ".join(ALLOWED_EXTENSIONS)}'})
        
        # Secure filename and save
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Extract text
        extracted_text = extract_text_from_file(filepath, filename)
        
        # Clean up - delete temporary file
        try:
            os.remove(filepath)
        except Exception as e:
            print(f"Error deleting temp file: {e}")
        
        if extracted_text and len(extracted_text) > 10:
            word_count = len(extracted_text.split())
            return jsonify({
                'success': True, 
                'text': extracted_text,
                'word_count': word_count,
                'file_name': filename,
                'file_type': filename.rsplit('.', 1)[1].lower()
            })
        else:
            return jsonify({'success': False, 'error': 'Could not extract text from file. Make sure the file contains readable text.'})
    
    except Exception as e:
        print(f"Upload error: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'})

@app.route('/api/ocr-status', methods=['GET'])
def ocr_status():
    """Check if OCR is configured"""
    try:
        # Test if pytesseract is working
        test_text = pytesseract.image_to_string(Image.new('RGB', (100, 100), color='white'))
        return jsonify({'available': True, 'message': 'OCR is configured and ready'})
    except Exception as e:
        return jsonify({'available': False, 'message': 'OCR not configured. Install Tesseract for image text extraction.'})

@app.route('/api/register', methods=['POST'])
def register():
    """User registration"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        name = data.get('name', email.split('@')[0] if '@' in email else 'User')
        
        if not email or not password:
            return jsonify({'success': False, 'error': 'Email and password required'})
        
        if len(password) < 6:
            return jsonify({'success': False, 'error': 'Password must be at least 6 characters'})
        
        users_file = 'users.json'
        users = {}
        if os.path.exists(users_file):
            with open(users_file, 'r') as f:
                users = json.load(f)
        
        if email in users:
            return jsonify({'success': False, 'error': 'User already exists'})
        
        users[email] = {
            'password': hashlib.md5(password.encode()).hexdigest(),
            'name': name,
            'email': email,
            'created_at': datetime.now().isoformat(),
            'plan': 'free',
            'checks_used': 0
        }
        
        with open(users_file, 'w') as f:
            json.dump(users, f, indent=2)
        
        return jsonify({'success': True, 'message': 'Registration successful'})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/login', methods=['POST'])
def login():
    """User login"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        
        if not email or not password:
            return jsonify({'success': False, 'error': 'Email and password required'})
        
        users_file = 'users.json'
        users = {}
        if os.path.exists(users_file):
            with open(users_file, 'r') as f:
                users = json.load(f)
        
        if email in users and users[email]['password'] == hashlib.md5(password.encode()).hexdigest():
            session.permanent = True
            session['user'] = {
                'email': email,
                'name': users[email]['name'],
                'plan': users[email].get('plan', 'free')
            }
            return jsonify({'success': True, 'user': session['user']})
        
        return jsonify({'success': False, 'error': 'Invalid credentials'})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/logout', methods=['POST'])
def logout():
    """User logout"""
    session.pop('user', None)
    return jsonify({'success': True})

@app.route('/api/check-auth', methods=['GET'])
def check_auth():
    """Check if user is authenticated"""
    if 'user' in session:
        return jsonify({'logged_in': True, 'user': session['user']})
    return jsonify({'logged_in': False})

@app.route('/api/check', methods=['POST'])
def check():
    """Check plagiarism between two texts"""
    try:
        if 'user' not in session:
            return jsonify({'success': False, 'error': 'Please login first'})
        
        data = request.get_json()
        text1 = data.get('text1', '').strip()
        text2 = data.get('text2', '').strip()
        
        if not text1 or not text2:
            return jsonify({'success': False, 'error': 'Please enter both texts'})
        
        if len(text1) < 20 or len(text2) < 20:
            return jsonify({'success': False, 'error': 'Please enter at least 20 characters in each text for accurate analysis'})
        
        # Update user check count
        users_file = 'users.json'
        if os.path.exists(users_file):
            with open(users_file, 'r') as f:
                users = json.load(f)
            email = session['user']['email']
            if email in users:
                users[email]['checks_used'] = users[email].get('checks_used', 0) + 1
                with open(users_file, 'w') as f:
                    json.dump(users, f, indent=2)
        
        result = check_plagiarism(text1, text2)
        return jsonify(result)
    
    except Exception as e:
        print(f"Check error: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Analysis error: {str(e)}'})

# Import re for text cleaning
import re

if __name__ == '__main__':
    print("=" * 50)
    print("🔍 PlagioScan Server Starting...")
    print("=" * 50)
    print(f"📍 Server URL: http://127.0.0.1:5000")
    print(f"📁 Upload folder: {app.config['UPLOAD_FOLDER']}")
    print(f"✅ Allowed extensions: {', '.join(ALLOWED_EXTENSIONS)}")
    print("=" * 50)
    print("🚀 Press CTRL+C to stop the server")
    print("=" * 50)
    app.run(debug=True, port=5000, host='127.0.0.1')